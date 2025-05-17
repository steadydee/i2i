"""
DocxRender — Supabase Storage helper
===================================

• Downloads “templates/<template_id>.docx”
• Replaces every {{merge_field}} (even when Word split it across runs)
  in body text, tables, headers & footers.
• Uploads result to “documents/<template_id>/<uuid>.docx”
• Returns {"ui_event":"download_link", "url": …}

Environment vars
----------------
SUPABASE_DOC_BUCKET   – output bucket name (default: documents)
URL_EXPIRY_MIN        – signed-URL lifetime in minutes (default: 120)
"""
from __future__ import annotations

import io, os, uuid
from typing import Dict, Any

import docx                              # pip install python-docx
from backend.db import sb                # Supabase client

TEMPLATE_BUCKET = "templates"
OUTPUT_BUCKET   = os.getenv("SUPABASE_DOC_BUCKET", "documents")
URL_EXPIRY_SEC  = int(os.getenv("URL_EXPIRY_MIN", "120")) * 60


# ────────── storage helpers ───────────────────────────────────────────
def _download(bucket: str, path: str) -> bytes:
    """Handle storage3 DownloadFileResponse as well as raw bytes."""
    obj = sb.storage.from_(bucket).download(path)
    return obj.file if hasattr(obj, "file") else obj


def _upload(bucket: str, key: str, blob: bytes) -> str:
    store = sb.storage.from_(bucket)
    store.upload(
        key,
        blob,
        {
            "content-type":
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        },
    )
    signed = store.create_signed_url(key, URL_EXPIRY_SEC)

    if isinstance(signed, str):           # very old client
        return signed
    if isinstance(signed, dict):          # storage3 ≥0.6
        return signed.get("signedURL") or signed.get("signed_url")
    if hasattr(signed, "signed_url"):     # storage3 0.5
        return signed.signed_url
    raise RuntimeError("Unknown signed-URL response shape")


# ────────── run-splitting safe replacement ───────────────────────────
def _replace_in_runs(runs, mapping: Dict[str, Any]) -> None:
    """
    Merge all runs' text, replace {{field}} tokens, then push the
    modified string back into the original run sequence.
    """
    full = "".join(r.text for r in runs)
    replaced = full
    for k, v in mapping.items():
        replaced = replaced.replace(f"{{{{{k}}}}}", str(v))

    if replaced == full:      # nothing changed
        return

    runs[0].text = replaced
    for r in runs[1:]:
        r.text = ""           # collapse extra runs


# ────────── main class (LangChain Runnable-friendly) ─────────────────
class DocxRender:
    def __init__(self, template_id: str):
        # allow caller to specify with / without .docx
        self.template_id = template_id.rstrip(".docx")

    # LangChain expects .invoke(input_dict) → output_dict
    def invoke(self, inputs: Dict[str, Any], **_) -> Dict[str, Any]:
        tpl_path = f"{self.template_id}.docx"
        tpl_blob = _download(TEMPLATE_BUCKET, tpl_path)

        doc = docx.Document(io.BytesIO(tpl_blob))

        # replace in body paragraphs + headers & footers
        parts = [doc] + [sect.header for sect in doc.sections] + \
                [sect.footer for sect in doc.sections]
        for part in parts:
            for para in part.paragraphs:
                _replace_in_runs(para.runs, inputs)

        # replace inside tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_runs(para.runs, inputs)

        # save to buffer
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)

        key = f"{self.template_id}/{uuid.uuid4()}.docx"
        url = _upload(OUTPUT_BUCKET, key, out.read())

        return {"ui_event": "download_link", "url": url}
